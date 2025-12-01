# yh_rag_cloud_api/parsing_utils.py

import os
import re
import json
import pandas as pd
import cv2
import numpy as np
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
from typing import List, Dict, Any, Optional
from pathlib import Path
import docx
import traceback
import pikepdf  # <-- Import for rotation fix
import csv
import base64

from google.api_core.client_options import ClientOptions

from .settings import settings
from google.oauth2 import service_account
from googleapiclient.discovery import build, HttpError, Resource
from google.cloud import vision, storage
# --- FIXED GC IMPORTS ---
from google.cloud import documentai
from google.cloud.documentai_v1 import types as documentai_types
from google.cloud.exceptions import NotFound
# --- END FIX ---

# --- NEW IMPORT: Explicitly import Vision types for shared GCS structures ---
from google.cloud.vision_v1.types import GcsSource as VisionGcsSource
from google.cloud.vision_v1.types import GcsDestination as VisionGcsDestination

# --- END NEW IMPORT ---


# ---
# --- General Text/OCR/File Utilities ---
# ---

print(f"DEBUG: GOOGLE_APPLICATION_CREDENTIALS set to: {os.getenv('GOOGLE_APPLICATION_CREDENTIALS')}")


def preprocess_image_for_ocr(pil_image: Image.Image) -> np.ndarray:
    """
    Converts an in-memory PIL Image to an OpenCV-compatible NumPy array
    and applies preprocessing for better OCR.
    """
    try:
        open_cv_image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
        gray = cv2.cvtColor(open_cv_image, cv2.COLOR_BGR2GRAY)
        binary = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 11, 2
        )
        denoised = cv2.medianBlur(binary, 3)
        return denoised
    except Exception as e:
        print(f"      ERROR during image preprocessing: {e}")
        return np.array(pil_image)


def extract_text_with_tables(section_pdf_path: str) -> str:
    """
    Extracts text using OCR (pytesseract) for every page,
    now with a pre-processing step to fix page rotation. (v20)
    """
    print(f"  Extracting text from {section_pdf_path} (FORCE OCR v20 - with rotation fix)...")
    full_text_ocr = ""

    doc = None
    pdf_bytes = None

    try:
        # --- NEW: Rotation Correction using pikepdf ---
        pdf_doc = pikepdf.Pdf.open(section_pdf_path)
        rotated = False

        for i, page in enumerate(pdf_doc.pages):
            if page.get("/Rotate") and int(page.get("/Rotate")) != 0:
                print(f"    Fixing existing rotation on page {i + 1}...")
                page["/Rotate"] = 0
                rotated = True

        if rotated:
            print("    Saving de-rotated PDF to memory buffer...")
            in_memory_pdf = io.BytesIO()
            pdf_doc.save(in_memory_pdf)
            in_memory_pdf.seek(0)
            pdf_bytes = in_memory_pdf.read()
            pdf_doc.close()
            # Open the *corrected* PDF from memory
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        else:
            print("    No rotation detected. Opening file directly.")
            pdf_doc.close()
            # Open the original file
            doc = fitz.open(section_pdf_path)
        # --- END: Rotation Correction ---

        # --- Now, proceed with the original Tesseract OCR logic ---
        for page_num, page in enumerate(doc):
            try:
                pix = page.get_pixmap(dpi=300)
                img_bytes = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_bytes))

                processed_img_array = preprocess_image_for_ocr(img)

                custom_config = r'--oem 3 --psm 3'
                ocr_text = pytesseract.image_to_string(processed_img_array, config=custom_config)

                full_text_ocr += ocr_text.strip() + "\n\n--- Page Break ---\n"
            except Exception as ocr_page_err:
                print(f"      ERROR during OCR for page {page_num + 1}: {ocr_page_err}")
                full_text_ocr += f"\n\n--- Error on Page {page_num + 1} ---\n\n"

        doc.close()
        return "[OCR Extracted Text]\n\n" + full_text_ocr

    except Exception as ocr_err:
        print(f"    FATAL Error during OCR processing: {ocr_err}")
        if doc:
            doc.close()
        return f"Error: Could not extract text from {section_pdf_path}"


def _get_storage_client():
    """Helper to build and return an authenticated Google Cloud Storage client."""
    # This will use the same GOOGLE_APPLICATION_CREDENTIALS as other clients
    if settings.google_application_credentials:
        return storage.Client.from_service_account_json(settings.google_application_credentials)
    return storage.Client()


def _get_documentai_client() -> Resource:
    """Helper to build and return an authenticated Document AI client."""
    # The Document AI API uses a different client discovery pattern
    from google.api_core.client_options import ClientOptions

    location = settings.document_ai_location
    # --- THIS IS THE CORRECT LINE ---
    opts = ClientOptions(api_endpoint=f"https://{location}-documentai.googleapis.com")

    # This will use GOOGLE_APPLICATION_CREDENTIALS if set, or ADC otherwise
    return build("documentai", "v1", client_options=opts)


def extract_text_with_document_ai(file_bytes: bytes, mime_type: str) -> str:
    """
    Extracts text from a document using Google Cloud Document AI.
    Uses ASYNCHRONOUS BATCH PROCESSING for reliable handling of large documents (up to 2,000 pages).
    """
    print(f"  [Document AI] Processing file ({mime_type}) using ASYNCHRONOUS BATCH PROCESS...")

    # Use the synchronous client for the initial operation submission
    client = documentai.DocumentProcessorServiceClient(
        client_options=ClientOptions(
            api_endpoint=f"{settings.document_ai_location}-documentai.googleapis.com"
        )
    )
    storage_client = _get_storage_client()

    # --- FIXED: Use the correct Document AI types ---
    # These are the proper types for the current Document AI library
    from google.cloud.documentai_v1.types import document_processor_service

    gcs_bucket_name = settings.gcs_bucket_name
    gcs_source_uri = ""
    gcs_destination_uri = ""

    # The full resource name of the processor
    processor_name = (
        f"projects/{settings.firebase_project_id}/locations/{settings.document_ai_location}"
        f"/processors/{settings.document_ai_processor_id}"
    )

    try:
        # 1. Upload file to GCS
        import uuid
        temp_filename = f"documentai_temp/{uuid.uuid4()}_{Path(uuid.uuid4().hex[:8]).name}.{'pdf' if mime_type == 'application/pdf' else 'docx'}"
        blob = storage_client.bucket(gcs_bucket_name).blob(temp_filename)
        blob.upload_from_string(file_bytes, content_type=mime_type)
        gcs_source_uri = f"gs://{gcs_bucket_name}/{temp_filename}"
        print(f"  [Document AI] Uploaded temp file to {gcs_source_uri}")

        # 2. Configure and submit async batch process request

        # Input Config - Use the proper GcsDocument type
        gcs_source = documentai_types.GcsDocuments(
            documents=[documentai_types.GcsDocument(gcs_uri=gcs_source_uri, mime_type=mime_type)]
        )
        input_config = documentai_types.BatchDocumentsInputConfig(gcs_documents=gcs_source)

        # Output Config
        gcs_destination_prefix = f"documentai_temp_output/{uuid.uuid4()}_output/"
        gcs_destination_uri = f"gs://{gcs_bucket_name}/{gcs_destination_prefix}"
        output_config = documentai_types.DocumentOutputConfig(
            gcs_output_config=documentai_types.DocumentOutputConfig.GcsOutputConfig(
                gcs_uri=gcs_destination_uri
            )
        )

        # Request - Use the proper BatchProcessRequest
        request = document_processor_service.BatchProcessRequest(
            name=processor_name,
            input_documents=input_config,
            document_output_config=output_config
        )

        print("  [Document AI] Submitting async batch process request (Timeout: 420s)...")
        operation = client.batch_process_documents(request=request)

        # Wait for the operation to complete
        operation.result(timeout=420)
        print("  [Document AI] Async operation finished.")

        # 3. Download and parse results from GCS
        # The output is a JSON file for each input document, or multiple files if the batch_size was > 1
        blob_list = storage_client.list_blobs(gcs_bucket_name, prefix=gcs_destination_prefix)
        full_text = ""

        for output_blob in blob_list:
            if output_blob.name.endswith('.json'):
                json_string = output_blob.download_as_string()
                response = json.loads(json_string)
                # The structure contains a 'text' field at the top level of the JSON document
                full_text += response.get('text', '')

        print(f"  [Document AI] Successfully extracted {len(full_text)} characters.")
        return full_text

    except Exception as e:
        # Re-print traceback for better debug info
        traceback.print_exc()
        print(f"    FATAL Error during Document AI batch processing: {e}")
        return f"Error: Could not process document with Document AI Batch API: {e}"
    finally:
        # 4. Clean up GCS files
        bucket = storage_client.bucket(gcs_bucket_name)
        if gcs_source_uri:
            try:
                source_blob_name = gcs_source_uri.replace(f"gs://{gcs_bucket_name}/", "")
                bucket.blob(source_blob_name).delete()
                print(f"  [Document AI] Cleaned up source blob: {source_blob_name}")
            except Exception as cleanup_error:
                print(f"    WARN: Failed to clean up source GCS blob {gcs_source_uri}: {cleanup_error}")
        if gcs_destination_uri:
            try:
                dest_prefix = gcs_destination_uri.replace(f"gs://{gcs_bucket_name}/", "")
                blobs_to_delete = list(bucket.list_blobs(prefix=dest_prefix))
                if blobs_to_delete:
                    bucket.delete_blobs(blobs_to_delete)
                    print(f"  [Document AI] Cleaned up {len(blobs_to_delete)} destination blobs.")
            except Exception as cleanup_error:
                print(f"    WARN: Failed to clean up destination GCS blobs at {gcs_destination_uri}: {cleanup_error}")


def extract_text_with_google_vision(file_bytes: bytes, mime_type: str = "application/pdf") -> str:
    """
    Extracts text from a document using Google Cloud Vision AI's Document Text Detection.
    Uses asynchronous batch processing for multi-page documents like PDFs.
    """
    print("  [Google Cloud Vision] Extracting text with Document AI...")
    client = vision.ImageAnnotatorClient()
    storage_client = _get_storage_client()
    gcs_bucket_name = settings.gcs_bucket_name
    gcs_source_uri = ""
    gcs_destination_uri = ""

    try:
        # 1. Upload file to GCS
        try:
            bucket = storage_client.get_bucket(gcs_bucket_name)
        except NotFound:
            print(f"    WARN: GCS bucket '{gcs_bucket_name}' not found. Attempting to create it...")
            # Buckets associated with Firebase projects are typically multi-region 'US'.
            # If your project is based elsewhere, you might need to specify a location.
            # See: https://cloud.google.com/storage/docs/locations
            bucket = storage_client.create_bucket(gcs_bucket_name, location="US")
            print(f"    SUCCESS: Created GCS bucket '{gcs_bucket_name}'.")

        import uuid
        temp_filename = f"vision_temp/{uuid.uuid4()}.pdf"
        blob = bucket.blob(temp_filename)
        blob.upload_from_string(file_bytes, content_type=mime_type)
        gcs_source_uri = f"gs://{gcs_bucket_name}/{temp_filename}"
        print(f"  [Google Cloud Vision] Uploaded temp file to {gcs_source_uri}")

        # 2. Configure and submit async request
        gcs_source = vision.GcsSource(uri=gcs_source_uri)
        input_config = vision.InputConfig(gcs_source=gcs_source, mime_type=mime_type)

        gcs_destination_prefix = f"vision_temp/{uuid.uuid4()}_output/"
        gcs_destination_uri = f"gs://{gcs_bucket_name}/{gcs_destination_prefix}"
        gcs_destination = vision.GcsDestination(uri=gcs_destination_uri)
        output_config = vision.OutputConfig(gcs_destination=gcs_destination, batch_size=100)

        feature = vision.Feature(type_=vision.Feature.Type.DOCUMENT_TEXT_DETECTION)
        async_request = vision.AsyncAnnotateFileRequest(
            features=[feature], input_config=input_config, output_config=output_config
        )

        print("  [Google Cloud Vision] Submitting async batch annotation request...")
        operation = client.async_batch_annotate_files(requests=[async_request])
        operation.result(timeout=420)  # Wait for the operation to complete
        print("  [Google Cloud Vision] Async operation finished.")

        # 3. Download and parse results from GCS
        blob_list = storage_client.list_blobs(gcs_bucket_name, prefix=gcs_destination_prefix)
        full_text = ""
        for output_blob in blob_list:
            json_string = output_blob.download_as_string()
            response = json.loads(json_string)
            for page_response in response.get('responses', []):
                annotation = page_response.get('fullTextAnnotation', {})
                full_text += annotation.get('text', '')

        print(f"  [Google Cloud Vision] Successfully extracted {len(full_text)} characters.")
        return full_text

    except Exception as e:
        print(f"    FATAL Error during Google Vision processing: {e}")
        return f"Error: Could not process document with Google Vision AI: {e}"
    finally:
        # 4. Clean up GCS files
        if gcs_source_uri:
            try:
                source_blob_name = gcs_source_uri.replace(f"gs://{gcs_bucket_name}/", "")
                bucket = storage_client.get_bucket(gcs_bucket_name)
                bucket.blob(source_blob_name).delete()
                print(f"  [Google Cloud Vision] Cleaned up source blob: {source_blob_name}")
            except Exception as cleanup_error:
                print(f"    WARN: Failed to clean up source GCS blob {gcs_source_uri}: {cleanup_error}")
        if gcs_destination_uri:
            try:
                dest_prefix = gcs_destination_uri.replace(f"gs://{gcs_bucket_name}/", "")
                bucket = storage_client.get_bucket(gcs_bucket_name)
                blobs_to_delete = list(bucket.list_blobs(prefix=dest_prefix))
                if blobs_to_delete:
                    bucket.delete_blobs(blobs_to_delete)
                    print(f"  [Google Cloud Vision] Cleaned up {len(blobs_to_delete)} destination blobs.")
            except Exception as cleanup_error:
                print(f"    WARN: Failed to clean up destination GCS blobs at {gcs_destination_uri}: {cleanup_error}")


def extract_text_from_docx(docx_path: str) -> str:
    """Extracts plain text content from a .docx file, including tables."""
    print(f"  Extracting text from DOCX: {docx_path}...")
    try:
        doc = docx.Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                # Join cells with a tab, join rows with a newline
                full_text.append('\t'.join(row_text))

        print(f"  Successfully extracted ~{len(full_text)} paragraphs/rows from DOCX.")
        return '\n'.join(full_text)
    except Exception as e:
        print(f"  ERROR: Failed to extract text from DOCX {docx_path}: {e}")
        return f"Error extracting DOCX: {e}"


def _create_text_chunks(text: str, chunk_size: int = 512, chunk_overlap: int = 100) -> List[str]:
    """
    Splits text into fixed-size overlapping chunks.
    """
    if not text: return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - chunk_overlap
    return chunks


def find_header_row(excel_path: str, keywords: List[str]) -> Optional[int]:
    """Tries to find the row index containing most of the keywords."""
    try:
        # Read without header first to inspect
        df_no_header = pd.read_excel(excel_path, sheet_name=0, header=None)
        max_matches = 0
        header_row_index = None
        # Check the first few rows (e.g., up to row 10)
        for i, row in df_no_header.head(10).iterrows():
            row_text = ' '.join(str(cell).lower() for cell in row if pd.notna(cell))
            matches = sum(1 for keyword in keywords if keyword in row_text)
            if matches > max_matches and matches >= len(keywords) // 2:  # Require at least half the keywords
                max_matches = matches
                header_row_index = i
        return header_row_index
    except Exception as e:
        print(f"  ERROR finding header row: {e}")
        return None


def excel_to_csv(excel_path: Path, sheet_name: Any = 0) -> str:
    """
    Converts a specific sheet from an Excel file to a temporary CSV file.

    Args:
        excel_path (Path): Path to the input .xlsx file.
        sheet_name (Any): The sheet name (str) or index (int) to convert.

    Returns:
        str: The file path to the created temporary CSV file.
    """
    try:
        # Read the specified sheet, without assuming a header
        df = pd.read_excel(excel_path, sheet_name=sheet_name, header=None)

        # Create a temp CSV path in the same directory
        # Use a descriptive name to avoid collisions
        csv_path = excel_path.with_name(f"{excel_path.stem}_sheet_{sheet_name}.csv")

        # Save to CSV without pandas index or header
        # Use csv.QUOTE_ALL to ensure all fields are quoted, handling commas
        df.to_csv(
            csv_path,
            index=False,
            header=False,
            quoting=csv.QUOTE_ALL
        )
        return str(csv_path)
    except Exception as e:
        print(f"  ERROR: Failed to convert Excel sheet '{sheet_name}' to CSV: {e}")
        print(traceback.format_exc())
        raise  # Re-raise the exception to be caught by the endpoint


def extract_text_with_document_ai_sync(file_bytes: bytes, mime_type: str) -> str:
    """
    Extract text from documents using Document AI SYNCHRONOUS processing.
    More reliable for single files than batch processing.
    """
    try:
        from google.cloud import documentai

        # Use your existing Document AI configuration
        project_id = "fluyx-yh-kb"
        location = "us"
        processor_id = "a5e8c8b3b57e7b1c"  # Your form parser processor ID

        client = documentai.DocumentProcessorServiceClient()
        processor_name = client.processor_path(project_id, location, processor_id)

        # For synchronous processing, we process directly without Cloud Storage
        raw_document = documentai.RawDocument(content=file_bytes, mime_type=mime_type)

        request = documentai.ProcessRequest(name=processor_name, raw_document=raw_document)
        result = client.process_document(request=request)

        document = result.document
        return document.text

    except Exception as e:
        return f"Error: Document AI sync processing failed: {str(e)}"